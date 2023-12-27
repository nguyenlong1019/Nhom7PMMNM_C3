from tkinter import *
from tkinter import ttk, filedialog, messagebox
from PIL import Image, ImageTk 
# import numpy as np
import pandas as pd 
import matplotlib.pyplot as plt
from matplotlib.figure import Figure 
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg 
import re 
import smtplib 
from email.message import EmailMessage
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm
import os 



class App:
    def __init__(self):
        self.root = Tk()
        self.root.title("Analysis Application")
        self.root.geometry('1400x700')
        self.root.resizable(0, 0)
        self.make_center()

        #---------------------------------
        # Developed By Nguyen Van Long 
        #---------------------------------
        self.analysis_data = None
        
        self.raw_data = None
        self.clean_data = None
        self.cols_visual = None

        # quản lý các biểu đồ
        # self.figure1 = None 
        # self.figure2 = None 
        # self.figure3 = None 
        # self.figure4 = None
        self.count = None
        self.data = None

        self.c11 = StringVar()
        self.c12 = StringVar()
        self.c21 = StringVar()
        self.c22 = StringVar()
        self.c31 = StringVar()
        self.c32 = StringVar()
        self.c41 = StringVar()
        self.c42 = StringVar()

        self.figure1 = [self.c11.get(), self.c12.get()]
        self.figure2 = [self.c21.get(), self.c22.get()]
        self.figure3 = [self.c31.get(), self.c32.get()]
        self.figure4 = [self.c41.get(), self.c42.get()]

        # style for all interfaces in project
        self.style = ttk.Style()

        # style name, options
        # Frame 1 Left
        self.style.configure(style="F1L.TFrame", background="#fff")
        # Frame 1 Right, Raw Data
        self.style.configure(style="F1R.RD.TFrame", background='#FFF7D4')

        # Frame 1 Left
        self.f1_left = ttk.Frame(self.root, style="F1L.TFrame")
        self.f1_left.place(x=0, y=0, width=400, height=700)
        
        # Button dùng để nhập file
        self.import_btn = ttk.Button(self.f1_left, text="Import Data", command=self.import_data)
        self.import_btn.place(x=10, y=10)
        
        # Button dùng để làm sạch data
        self.clean_data_btn = ttk.Button(self.f1_left, text="Clean Data", command=self.clean_data_f)
        self.clean_data_btn.place(x=10, y=50)

        self.style.configure(style="CC.TFrame", background='#F3F8FF')

        # Frame này để thực hiện khi người dùng muốn trực quan hóa dữ liệu,
        # Người dùng sẽ chọn các trường để phân tích
        self.choose_chart_frm = ttk.Frame(self.f1_left, style='CC.TFrame')
        self.choose_chart_frm.place(x=10, y=100, width=380, height=500)

        self.number_of_chart = ttk.Scale(self.choose_chart_frm, orient="horizontal", from_=1, to=4, command=self.on_scale_change)
        self.number_of_chart.place(x=5,y=5)
        self.number_of_chart_label = ttk.Label(self.choose_chart_frm, text="Charts: ")
        self.number_of_chart_label.place(x=5,y=35)
        # self.create_form_btn = ttk.Button(self.choose_chart_frm, text="Create")
        # self.create_form_btn.place(x=100, y=35)

        self.wrapper_frame = ttk.Frame(self.choose_chart_frm)
        self.wrapper_frame.place(x=5, y=70, width=360, height=400)
        
        
        
        #######################################################

        # Button để quit app
        self.quit_btn = ttk.Button(self.f1_left, text="Quit", command=quit)
        self.quit_btn.place(x=10, y=660)
        
        # Button để gửi phản hồi về ứng dụng đến nhà phát triển
        self.feedback_btn = ttk.Button(self.f1_left, text="Feedback", command=self.feedback_layot)
        self.feedback_btn.place(x=100, y=660)

        self.separator_left = ttk.Separator(self.f1_left, orient='vertical')
        self.separator_left.place(x=395, y=0, width=5, height=700)

        # Frame 2 Right
        self.f1_right = ttk.Frame(self.root)
        self.f1_right.place(x=400, y=0, width=1000, height=700)

        # Frame to contain the raw data table
        self.raw_data_frm = ttk.Frame(self.f1_right)
        self.raw_data_frm.place(x=0, y=0, width=1000, height=300)

        # selectmode = 'extended' enable choose multiple at a time
        self.raw_data_table = ttk.Treeview(self.raw_data_frm, selectmode='extended')
        self.raw_data_table.place(x=0, y=0, width=950, height=250)

        # Vertical Scrollbar
        self.v_scrollbar_rdt = ttk.Scrollbar(self.raw_data_frm, orient='vertical', command=self.raw_data_table.yview)
        self.raw_data_table.configure(yscrollcommand=self.v_scrollbar_rdt.set)
        self.v_scrollbar_rdt.place(x=950, y=0, width=50, height=250)

        # Horizontal Scrollbar
        self.h_scrollbar_rdt = ttk.Scrollbar(self.raw_data_frm, orient='horizontal', command=self.raw_data_table.xview)
        self.raw_data_table.configure(xscrollcommand=self.h_scrollbar_rdt.set)
        self.h_scrollbar_rdt.place(x=0,y=250, width=950, height=30)

        self.separator_rdt = ttk.Separator(self.raw_data_frm, orient='horizontal')
        self.separator_rdt.place(x=0, y=290, width=1000, height=10)

        # define number of columns, max columns = 10
        self.raw_data_table["columns"] = ("1", "2", "3", "4", "5", "6", "7", "8", "9", "10")

        # define heading
        self.raw_data_table["show"] = 'headings'

        # "n": hướng bắc (căn về phía bên trên)
        # "s": hướng nam (căn về phía bên dưới
        # "e": hướng đông (căn về phía bên phải)
        # "w": hướng tây (căn về phía bên trái)
        # "ne": hướng đông bắc (căn về phía trên bên phải)
        # "nw": hướng tây bắc (căn về phía trên bên trái)
        # "se": hướng đông nam (căn về phía dưới bên phải)
        # "sw": hướng tây nam (căn về phía dưới bên trái)
        # "c" hoặc "center": căn chính giữa
        self.raw_data_table.column("1", anchor='c')
        self.raw_data_table.column("2", anchor='c')
        self.raw_data_table.column("3", anchor='c')
        self.raw_data_table.column("4", anchor='c')
        self.raw_data_table.column("5", anchor='c')
        self.raw_data_table.column("6", anchor='c')
        self.raw_data_table.column("7", anchor='c')
        self.raw_data_table.column("8", anchor='c')
        self.raw_data_table.column("9", anchor='c')
        self.raw_data_table.column("10", anchor='c')

        self.raw_data_table.heading("1", text="")
        self.raw_data_table.heading("2", text="")
        self.raw_data_table.heading("3", text="")
        self.raw_data_table.heading("4", text="")
        self.raw_data_table.heading("5", text="")
        self.raw_data_table.heading("6", text="")
        self.raw_data_table.heading("7", text="")
        self.raw_data_table.heading("8", text="")
        self.raw_data_table.heading("9", text="")
        self.raw_data_table.heading("10", text="")

        for i in range(100):
            self.raw_data_table.insert(parent="", index="end", iid=f"item{i}", values=("", "", "", "", "", "", "", "", "", ""))
            pass

        # Frame to contain the clean data table
        self.clean_data_frm = ttk.Frame(self.f1_right)
        self.clean_data_frm.place(x=0, y=300, width=1000, height=400)

        self.clean_data_table = ttk.Treeview(self.clean_data_frm, selectmode='extended')
        self.clean_data_table.place(x=0, y=0, width=950, height=300)

        # Vertical Scrollbar
        self.v_scrollbar_cdt = ttk.Scrollbar(self.clean_data_frm, orient='vertical', command=self.clean_data_table.yview)
        self.clean_data_table.configure(yscrollcommand=self.v_scrollbar_cdt.set)
        self.v_scrollbar_cdt.place(x=950, y=0, width=50, height=300)

        # Horizontal Scrollbar
        self.h_scrollbar_cdt = ttk.Scrollbar(self.clean_data_frm, orient='horizontal', command=self.clean_data_table.xview)
        self.clean_data_table.configure(xscrollcommand=self.h_scrollbar_cdt.set)
        self.h_scrollbar_cdt.place(x=0,y=300, width=950, height=20)

        # define number of columns, max columns = 10
        self.clean_data_table["columns"] = ("1", "2", "3", "4", "5", "6", "7", "8", "9", "10")

        # define heading
        self.clean_data_table["show"] = 'headings'

        self.clean_data_table.column("1", anchor='c')
        self.clean_data_table.column("2", anchor='c')
        self.clean_data_table.column("3", anchor='c')
        self.clean_data_table.column("4", anchor='c')
        self.clean_data_table.column("5", anchor='c')
        self.clean_data_table.column("6", anchor='c')
        self.clean_data_table.column("7", anchor='c')
        self.clean_data_table.column("8", anchor='c')
        self.clean_data_table.column("9", anchor='c')
        self.clean_data_table.column("10", anchor='c')

        self.clean_data_table.heading("1", text="")
        self.clean_data_table.heading("2", text="")
        self.clean_data_table.heading("3", text="")
        self.clean_data_table.heading("4", text="")
        self.clean_data_table.heading("5", text="")
        self.clean_data_table.heading("6", text="")
        self.clean_data_table.heading("7", text="")
        self.clean_data_table.heading("8", text="")
        self.clean_data_table.heading("9", text="")
        self.clean_data_table.heading("10", text="")

        for i in range(100):
            self.clean_data_table.insert(parent="", index="end", iid=f"item{i}", values=("", "", "", "", "", "", "", "", "", ""))
            pass

        # Button để chuyển sang trang trực quan hóa dữ liệu
        self.visualization_btn = ttk.Button(self.clean_data_frm, text="Visualization", command=self.switch_frames_next)
        self.visualization_btn.place(x=10, y=360)

        # Button để download clean data
        self.download_clean_data = ttk.Button(self.clean_data_frm, text="Download Clean Data", command=self.donwload_clean_data_f)
        # self.download_clean_data.place(x=100, y=360)


        ###########################################
        # FRAME 2
        ###########################################
        self.f2_left = ttk.Frame(self.root)
        # self.f2_left.place(x=0, y=0, width=400, height=700)
        self.f2_home_btn = ttk.Button(self.f2_left, text='Home', command=self.switch_frames_back)
        self.f2_home_btn.place(x=10, y=10)
        self.download_figure_btn = ttk.Button(self.f2_left, text="Download", command=self.save_img_to_file)
        self.download_figure_btn.place(x=50,y=200)
        self.figures = ttk.Combobox(self.f2_left, values=("Figure 1", "Figure 2", "Figure 3", "Figure 4"))
        self.figures.place(x=150, y=200)
        self.figures.current(0)
        self.f2_next_btn = ttk.Button(self.f2_left, text='Analysis', command=self.switch_frames_next)
        self.f2_next_btn.place(x=200, y=550)

        self.separator_left_2 = ttk.Separator(self.f2_left, orient='vertical')
        self.separator_left_2.place(x=395, y=0, width=5, height=700)

        # test = ttk.Style()
        # test.configure("C1.TFrame",background="yellow")
        # test.configure("C2.TFrame",background="red")
        # test.configure("C3.TFrame",background="blue")
        # test.configure("C4.TFrame",background="pink")

        self.f2_right = ttk.Frame(self.root)
        # self.f2_right.place(x=400, y=0, width=1000, height=700)
        # self.f2_lbl = ttk.Label(self.f2_right, text="F2")
        # self.f2_lbl.place(x=10,y=50)

        # tạo dữ liệu test hiển thị biểu đồ

        self.cv1 = ttk.Frame(self.f2_right)
        self.cv1.place(x=0,y=0,width=480, height=330)
        # self.cv1.configure(style="C1.TFrame")

        # categories = ['Mitsubishi', 'Vinfast Lux', 'VietNamHanoi', 'ChuongMy', 'DaiHoc', 'QQQQQQQ', 'G', 'H']
        # values = [25, 40, 30, 20, 70, 40, 55, 665]

        self.fig1 = Figure(figsize=(6,4), tight_layout=True, dpi=100)
        self.ax1 = self.fig1.add_subplot(111)
        # lấy dải màu
        

        # vẽ biểu đồ cột
        # self.ax1.barh(categories, values, color=colormap(values))
        # self.ax1.set_xlabel("Thể loại")
        # self.ax1.set_ylabel("Số lượng")
        # self.ax1.set_title("Top 10 số lượng theo thể loại")
        
        self.canvas1 = FigureCanvasTkAgg(self.fig1, master=self.cv1)
        self.canvas1.get_tk_widget().pack(expand=True, fill='both')


        self.cv2 = ttk.Frame(self.f2_right)
        self.cv2.place(x=520,y=0,width=480,height=330)

        # x_data = [1, 2, 3, 4, 5, 6, 7, 8,9,10,11,12,13,14,15,16,17,18]
        # y_data1 = [25, 40, 30, 20, 70, 40, 55, 22,34,54,43,65,76,76,45,34,44,65]
        # y_data2 = [15, 35, 20, 30, 50, 45, 60, 30,54,43,33,95,86,60,99,87,11,97]

        self.fig2 = Figure(figsize=(6,4), tight_layout=True, dpi=100)
        self.ax2 = self.fig2.add_subplot(111)

        # self.ax2.plot(x_data, y_data1, label="Dữ liệu 1", marker='o', linestyle='-', color='skyblue')
        # self.ax2.plot(x_data, y_data2, label="Dữ liệu 2", marker='s', linestyle='--', color='orange')

        # self.ax2.set_xlabel('Thể loại')
        # self.ax2.set_ylabel('Số lượng')
        # self.ax2.set_title('Data 1 and Data 2')
        # self.ax2.legend()

        self.canvas2 = FigureCanvasTkAgg(self.fig2, master=self.cv2)
        self.canvas2.get_tk_widget().pack(expand=True, fill='both')

        


        self.cv3 = ttk.Frame(self.f2_right)
        self.cv3.place(x=0,y=370,width=480,height=330)

        self.fig3 = Figure(figsize=(6,4), tight_layout=True, dpi=100)
        self.ax3 = self.fig3.add_subplot(111)

        self.canvas3 = FigureCanvasTkAgg(self.fig3, master=self.cv3)
        self.canvas3.get_tk_widget().pack(expand=True, fill='both')



        self.cv4 = ttk.Frame(self.f2_right)
        self.cv4.place(x=520,y=370,width=480,height=330)

        self.fig4 = Figure(figsize=(6,4), tight_layout=True, dpi=100)
        self.ax4 = self.fig4.add_subplot(111)

        self.canvas4 = FigureCanvasTkAgg(self.fig3, master=self.cv4)
        self.canvas4.get_tk_widget().pack(expand=True, fill='both')        




        ###########################################
        # FRAME 3
        ###########################################
        self.f3_left = ttk.Frame(self.root)
        # self.f3_left.place(x=0, y=0, width=400, height=700)
        self.f3_home_btn = ttk.Button(self.f3_left, text='Home', command=self.switch_frames_back)
        self.f3_home_btn.place(x=10, y=10)
        
        self.f3_export_word = ttk.Button(self.f3_left, text='Export Word', command=self.export_word)
        self.f3_export_word.place(x=200, y=550)

        self.res_lbl = ttk.Label(self.f3_left, text="Lưu thành công!")
        # self.res_lbl.place(x=200,y=650)

        self.separator_left_3 = ttk.Separator(self.f3_left, orient='vertical')
        self.separator_left_3.place(x=395, y=0, width=5, height=700)

        style_f3 = ttk.Style()
        style_f3.configure("F3R.TFrame", background='#ffffff')
        self.f3_right = ttk.Frame(self.root)
        self.f3_right.configure(style="F3R.TFrame")
        # self.f3_right.place(x=400, y=0, width=1000, height=700)
        # self.f3_lbl = ttk.Label(self.f3_right, text="F3")
        # self.f3_lbl.place(x=10,y=50)

        # header text style
        self.header_font = ("Times New Roman", 20, 'bold')
        self.text_font = ("Times New Roman", 14, 'normal')

        # canvas
        self.canvas_f3 = Canvas(self.f3_right, background='#fff',width=900, height=700)
        self.canvas_f3.pack(side=LEFT)

        # Scorllbar
        self.y_scrollbar = ttk.Scrollbar(self.f3_right,orient=VERTICAL,command=self.canvas_f3.yview)
        self.y_scrollbar.pack(side=RIGHT,fill=Y)

        self.canvas_f3.configure(yscrollcommand=self.y_scrollbar.set)

        # Content Frame
        self.content_frame = ttk.Frame(self.canvas_f3,width=900,height=2600,style="F3R.TFrame")
        self.canvas_f3.create_window((0,0),window=self.content_frame,anchor=NW)

        self.header_text = Label(self.content_frame,text='Báo cáo phân tích dữ liệu', anchor=CENTER, justify=CENTER, background='#fff')
        self.header_text.configure(font=self.header_font)
        self.header_text.place(x=450,y=20,anchor=CENTER)

        self.body_text = Label(self.content_frame, text='I. Thông tin chung',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
        self.body_text.place(x=50,y=100)



        # Viết code ở đây




        # self.body_text1 = Label(self.content_frame, text='Tên Môn: Lập trình Python',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
        # self.body_text1.place(x=50,y=130)

        # self.body_text2 = Label(self.content_frame, text='Mã môn: FE6051        Số tín chỉ: 3 (2,1,0)',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
        # self.body_text2.place(x=50,y=160)

        # self.body_text3 = Label(self.content_frame, text='Số lớp học phần: 9',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
        # self.body_text3.place(x=50,y=190)

        # self.body_text4 = Label(self.content_frame, text='Tổng số sinh viên: 700        Pass: 83.2%',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
        # self.body_text4.place(x=50,y=220)
        
        # self.body_text5 = Label(self.content_frame, text='II. Kết quả trực quan hóa dữ liệu',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
        # self.body_text5.place(x=50,y=250)

        # self.image1 = Label(self.content_frame, image=self.score_ratio, background='#fff', anchor=CENTER, justify=CENTER)
        # self.image1.place(x=450,y=450,anchor=CENTER)

        # self.image2 = Label(self.content_frame,image=self.test_ratio,background='#fff',anchor=CENTER,justify=CENTER)
        # self.image2.place(x=450,y=800,anchor=CENTER)

        # self.body_text6 = Label(self.content_frame, text='III. Kết luận',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
        # self.body_text6.place(x=50,y=1000)

        # self.body_text7 = Label(self.content_frame, text='.'*79,background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
        # self.body_text7.place(x=50,y=1030)

        # self.body_text8 = Label(self.content_frame, text='.'*79,background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
        # self.body_text8.place(x=50,y=1060)

        # end report
        # self.body_text9 = Label(self.content_frame, text='Ngày .... Tháng .... Năm .... ',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
        # self.body_text9.place(x=500,y=1100)

        # self.body_text9 = Label(self.content_frame, text='Ký tên',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
        # self.body_text9.place(x=575,y=1130)

        # self.content_frame.update_idletasks()
        # self.canvas_f3.config(scrollregion=self.canvas_f3.bbox("all"))

        self.content_frame.update_idletasks()
        self.canvas_f3.config(scrollregion=self.canvas_f3.bbox("all"))




        ######################################################
        # Send Email
        ######################################################
        self.feedback_frm = ttk.Frame(self.root)
        self.fb_to_home_btn = ttk.Button(self.feedback_frm, text="Home", command=self.feedback_layout_to_home)
        self.fb_to_home_btn.place(x=50,y=620)
        self.lbl_frm_3 = ttk.LabelFrame(self.feedback_frm, text="Gửi ý kiến của bạn đến nhóm phát triển")
        self.lbl_frm_3.grid(row=0, column=0, padx=50, pady=10)

        self.email_lbl = ttk.Label(self.lbl_frm_3, text="Nhập Email của bạn: ")
        self.email_lbl.grid(row=0, column=0, padx=10, pady=10)
        self.email_ent = ttk.Entry(self.lbl_frm_3, width=75)
        self.email_ent.grid(row=0, column=1, padx=10, pady=10)

        self.subject_mail = ttk.Label(self.lbl_frm_3, text="Nhập tiêu đề email: ")
        self.subject_mail.grid(row=1, column=0, padx=10, pady=10)
        self.subject_ent = ttk.Entry(self.lbl_frm_3, width=75)
        self.subject_ent.grid(row=1, column=1, padx=10, pady=10)

        self.body_mail = ttk.Label(self.lbl_frm_3, text="Nhập nội dung email:")
        self.body_mail.grid(row=2, column=0, padx=10, pady=10)

        self.body_mail_text = Text(self.lbl_frm_3, width=60)
        self.body_mail_text.grid(row=3, column=1, padx=10, pady=10)

        # self.attactment_btn = ttk.Button(self.lbl_frm_3, text='Đính kèm file')
        # self.attactment_btn.grid(row=4, column=1, padx=10, pady=10, sticky='w') # west: tây, căn trái
        self.submit_btn = ttk.Button(self.lbl_frm_3, text='Gửi', command=self.submit_form)
        self.submit_btn.grid(row=5, column=1, padx=10, pady=10, sticky='w')
        self.result_lbl = ttk.Label(self.lbl_frm_3, text='')
        self.result_lbl.grid(row=5, column=2, padx=10, pady=10)
        ######################################################
        # Send Email
        ######################################################






        self.current_frame_l = self.f1_left
        self.current_frame_r = self.f1_right


    def on_scale_change(self,value):
        self.clear_frame(self.wrapper_frame)
        print(f"Giá trị: {int(float(value))}")
        self.number_of_chart_label.configure(text=f"Charts: {int(float(value))}")
        self.count = int(float(value))
        var1 = [self.c11, self.c21, self.c31, self.c41]
        var2 = [self.c12, self.c22, self.c32, self.c42]
        if self.cols_visual is not None:
            if int(float(value)):
                for i in range(int(float(value))):
                    combobox1 = ttk.Combobox(self.wrapper_frame, textvariable=var1[i])
                    combobox1.grid(row=i, column=0, padx=5, pady=5)
                    combobox1['values'] = self.cols_visual
                    combobox1.current(i)
                    combobox2 = ttk.Combobox(self.wrapper_frame, textvariable=var2[i])
                    combobox2.grid(row=i, column=1, padx=5, pady=5)
                    combobox2['values'] = self.cols_visual
                    combobox2.current(i)
                self.clean_data_frm.after(1000, lambda: self.visualization_btn.place(x=10, y=360))

        else:
            messagebox.showerror(title="Error", message="Cần nhập dữ liệu để trực quan hóa dữ liệu")

    
    def clear_frame(self, frame):
        for widget in frame.winfo_children():
            widget.destroy()

        
    def import_data(self):
        import_file_path = filedialog.askopenfilename(defaultextension=".csv", filetypes=[("CSV files","*.csv")])
        if import_file_path:
            self.raw_data = pd.read_csv(import_file_path)
            cols_name = self.raw_data.columns

            self.raw_data_table.delete(*self.raw_data_table.get_children())
            for i in range(len(cols_name)):
                self.raw_data_table.heading(str(i+1), text=cols_name[i])
            
            for i, row in self.raw_data.iterrows():
                self.raw_data_table.insert(parent="", index="end", iid=f"item{i}", values=tuple(row))


    def clean_data_f(self):
        if self.raw_data is not None:
            self.clean_data = self.raw_data
            # self.clean_data = pd.DataFrame(self.raw_data)
            mean_values = self.clean_data.mean()
            self.clean_data.fillna(mean_values, inplace=True)  # điền các giá trị trống bằng mean() của cột đó
            
            # nếu các cột text quá dài thì xóa bớt và chỉ để 12 ký tự để hiển thị
            for column in self.clean_data.columns:
                if self.clean_data[column].dtype == 'O':
                    self.clean_data[column] = self.clean_data[column].str[:12]

            cols_name = self.clean_data.columns

            self.clean_data_table.delete(*self.clean_data_table.get_children())
            for i in range(len(cols_name)):
                self.clean_data_table.heading(str(i+1), text=cols_name[i])
            
            for i, row in self.clean_data.iterrows():
                self.clean_data_table.insert(parent="", index="end", iid=f"item{i}", values=tuple(row))

            self.clean_data_frm.after(1000, lambda: self.download_clean_data.place(x=100, y=360))
        
            # self.cols_visual = self.clean_data.select_dtypes(include=['number']).columns
            self.cols_visual = self.clean_data.columns.to_list()
        else:
            messagebox.showerror(title="Error", message="Dữ liệu thô không được phép trống!")


    def donwload_clean_data_f(self):
        download_clean_data_file_path = filedialog.asksaveasfilename(defaultextension='.csv', filetypes=[("CSV files", "*.csv")])
        if download_clean_data_file_path:
            self.clean_data.to_csv(download_clean_data_file_path, index=False)
            print("Download Clean Data Successfully!")


        #---------------------------------
        # Developed By Nguyen Van Long 
        #---------------------------------

    
    def feedback_layot(self):
        self.f1_left.place_forget()
        self.f1_right.place_forget()
        self.feedback_frm.place(x=0,y=0,width=1400,height=700)

        
    def feedback_layout_to_home(self):
        self.feedback_frm.place_forget()
        self.f1_left.place(x=0, y=0, width=400, height=700)
        self.f1_right.place(x=400, y=0, width=1000, height=700)

    
    def switch_frames_next(self):
        # chỉ từ frame 1 -> frame 2 -> frame 3
        if self.current_frame_l == self.f1_left and self.current_frame_r == self.f1_right:
            if self.count is not None:
                # khi chuyển sang visualization
                figure_list = [self.figure1, self.figure2, self.figure3, self.figure4]
                var1 = [self.c11, self.c21, self.c31, self.c41]
                var2 = [self.c12, self.c22, self.c32, self.c42]
                for i in range(self.count):
                    figure_list[i] = [var1[i].get(), var2[i].get()]
                
                self.figure1 = figure_list[0]
                self.figure2 = figure_list[1]
                self.figure3 = figure_list[2]
                self.figure4 = figure_list[3]
                
                # điều kiện chuyển tap
                # số biểu đồ không được phép none và 2 cột không được phép tất cả là text

                self.send_data()

                if len(self.data) > 0:
                    self.f1_left.place_forget()
                    self.f1_right.place_forget()
                    self.f2_left.place(x=0, y=0, width=400, height=700)
                    self.f2_right.place(x=400, y=0, width=1000, height=700)
                    self.current_frame_l = self.f2_left
                    self.current_frame_r = self.f2_right

                    self.plot_data()
                else:
                    messagebox.showerror(title="Error", message="Dữ liệu trực quan không hợp lệ!\n\
                        Các biểu đồ trực quan không được phép có 2 cột dạng text!!!")
            else:
                messagebox.showerror(title="Error", message="Số lượng biểu đồ cần phải chọn!")


        elif self.current_frame_l == self.f2_left and self.current_frame_r == self.f2_right:
            # phân tích dữ liệu ở đây
            self.auto_analysis()
            

            self.f2_left.place_forget()
            self.f2_right.place_forget()
            self.f3_left.place(x=0, y=0, width=400, height=700)
            self.f3_right.place(x=400, y=0, width=1000, height=700)
            self.current_frame_l = self.f3_left
            self.current_frame_r = self.f3_right

            self.draw_report()

    
    def auto_analysis(self):
        self.analysis_data = dict()
        num_rows, num_cols = self.clean_data.shape
        column_names = self.clean_data.columns.tolist()
        col_names_text = " ".join(column_names)
        recommend = f"Dựa trên dữ liệu phân tích của {num_rows} dòng dữ liệu và {num_cols} cột dữ liệu.\n"
        recommend += col_names_text

        
        
        numeric_cols = self.clean_data.select_dtypes(include='number')
        column_stats = pd.DataFrame({
            'Mean': numeric_cols.mean(),
            'Mode': numeric_cols.mode().iloc[0],  # Lấy mode với giả sử có nhiều mode
            'Median': numeric_cols.median(),
            'Std': numeric_cols.std(),
            'Min': numeric_cols.min(),
            'Max': numeric_cols.max()
        })
        # print(column_stats)
        self.analysis_data['recommend'] = recommend
        self.analysis_data['statistic'] = column_stats

        # print(self.analysis_data)
    

    def draw_report(self):
        if self.analysis_data is not None:
            # print(self.analysis_data['recommend'])
            # print(self.analysis_data['statistic'])
            self.body_text1 = Label(self.content_frame, text=self.analysis_data['recommend'],background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
            self.body_text1.place(x=50,y=130)

            print(self.analysis_data['statistic'])
            i = 0
            for index, row in self.analysis_data['statistic'].iterrows():
                print(row)
                text = f"{row.name}: mean = {row[1]}, mode = {row[2]}, median = {row[3]}, min = {row[4]}, max = {row[5]}"
                body_text = Label(self.content_frame, text=text, background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
                y_ = 190+30*i
                body_text.place(x=50, y=y_)
                i += 1
            length_ = len(self.analysis_data['statistic'])
            self.body_text5 = Label(self.content_frame, text='II. Kết quả trực quan hóa dữ liệu',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
            self.body_text5.place(x=50,y=190+30*length_)
            
            # self.body_text2 = Label(self.content_frame, text='Mã môn: FE6051        Số tín chỉ: 3 (2,1,0)',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
            # self.body_text2.place(x=50,y=160)

            # self.body_text3 = Label(self.content_frame, text='Số lớp học phần: 9',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
            # self.body_text3.place(x=50,y=190)

            # self.body_text4 = Label(self.content_frame, text='Tổng số sinh viên: 700        Pass: 83.2%',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
            # self.body_text4.place(x=50,y=220)
            
            img0 = 'pic0.png'
            img1 = 'pic1.png'
            img2 = 'pic2.png'
            img3 = 'pic3.png'
            length_1 = 190 + 30 * length_ + 200
            j = 0
            if os.path.exists(img0):
                self.pic0 = ImageTk.PhotoImage(Image.open(img0).resize((400,300)))
                self.image0 = Label(self.content_frame, image=self.pic0, background='#fff', anchor=CENTER, justify=CENTER)
                self.image0.place(x=450,y=length_1 + 350 * j,anchor=CENTER)
                j += 1
            
            if os.path.exists(img1):
                self.pic1 = ImageTk.PhotoImage(Image.open(img1).resize((400,300)))
                self.image1 = Label(self.content_frame, image=self.pic1, background='#fff', anchor=CENTER, justify=CENTER)
                self.image1.place(x=450,y=length_1 + 350 * j,anchor=CENTER)
                j += 1
            if os.path.exists(img2):
                self.pic2 = ImageTk.PhotoImage(Image.open(img2).resize((400,300)))
                self.image2 = Label(self.content_frame, image=self.pic2, background='#fff', anchor=CENTER, justify=CENTER)
                self.image2.place(x=450,y=length_1 + 350 * j,anchor=CENTER)
                j += 1
            if os.path.exists(img3):
                self.pic3 = ImageTk.PhotoImage(Image.open(img3).resize((400,300)))
                self.image3 = Label(self.content_frame, image=self.pic3, background='#fff', anchor=CENTER, justify=CENTER)
                self.image3.place(x=450,y=length_1 + 350 * j,anchor=CENTER)
                j += 1

            length_2 = length_1 + 350 * j + 30

            # m = 0
            self.body_text6 = Label(self.content_frame, text='III. Kết luận',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
            self.body_text6.place(x=50,y=length_2)

            self.body_text7 = Label(self.content_frame, text='.'*79,background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
            self.body_text7.place(x=50,y=length_2 + 30)

            self.body_text8 = Label(self.content_frame, text='.'*79,background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
            self.body_text8.place(x=50,y=length_2 + 60)

            self.body_text9 = Label(self.content_frame, text='Ngày .... Tháng .... Năm .... ',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
            self.body_text9.place(x=500,y=length_2 + 10)

            self.body_text9 = Label(self.content_frame, text='Ký tên',background='#fff', font=self.text_font, anchor=CENTER, justify=CENTER)
            self.body_text9.place(x=575,y=length_2 + 40)

            
    
    def send_data(self):
        '''Hàm lọc lấy các figure để vẽ biểu đồ hợp lệ
        
        Các biểu đồ hợp lệ: không có 2 cột vẽ trùng nhau, 2 biểu đồ không được lặp lại cột
        Ví dụ: [['A', 'B'], ['B', 'A'], ['A', 'A'], ['C', 'D']] => [['A','B'], ['C','D']]
        '''
        data_ = []
        if self.figure1 is not None:
            data_.append(self.figure1)
        if self.figure2 is not None:
            data_.append(self.figure2)
        if self.figure3 is not None:
            data_.append(self.figure3)
        if self.figure4 is not None:
            data_.append(self.figure4)
        
        seen = set()
        unique_data = []
        for item in data_:
            # Sắp xếp các phần tử trong mỗi sublist để ['A', 'B'] và ['B', 'A'] được coi là giống nhau
            sorted_item = tuple(sorted(item))

            # Nếu phần tử chưa xuất hiện, thêm vào mảng mới và set
            if sorted_item not in seen:
                seen.add(sorted_item)
                unique_data.append(list(sorted_item))

        unique_data = [row for row in unique_data if row[0] != row[1]]
        unique_data = [row for row in unique_data if self.clean_data[row[0]].dtype != 'object' or self.clean_data[row[1]].dtype != 'object']
        # sau khi lọc như này rồi còn phải kiểm tra các cột dữ liệu bên trong đó
        # for row in unique_data:
        #     type_column1 = df[row[0]].dtype
        #     type_column2 = df[row[1]].dtype
        #     if type_column1 == 'object' and type_column2 == 'object':
        #         pass
        self.data = unique_data
        print(self.data)


    def plot_data(self):
        if len(self.data) > 0:
            var_axs = [self.ax1, self.ax2, self.ax3, self.ax4]
            var_canvas = [self.canvas1, self.canvas2, self.canvas3, self.canvas4]
            var_fig = [self.fig1, self.fig2, self.fig3, self.fig4]
            for i, item in enumerate(self.data):
                # print(item)
                if (self.clean_data[item[0]].dtype == 'object') or (self.clean_data[item[1]].dtype == 'object'):
                    # vẽ biểu đồ cột
                    colormap = plt.get_cmap('viridis')
                    if self.clean_data[item[0]].dtype == 'object':
                        # biểu đồ cột thì sẽ lấy 5 cột max và 5 cột min của dữ liệu số
                        df_sorted = self.clean_data.sort_values(by=item[1])
                        df = pd.concat([df_sorted.head(5), df_sorted.tail(5)])
                        # print(df) vì sao lại chỉ có 9 hoặc 8 rồi, vì nó trùng giá trị nên trùng vị trí

                        text = df[item[0]]
                        values = df[item[1]]
                        var_axs[i].barh(text, values, color=colormap(values))
                        var_axs[i].set_xlabel(item[0])
                        var_axs[i].set_ylabel(item[1])
                        var_axs[i].set_title(f"Top 10")
                        var_canvas[i].draw()

                        var_fig[i].savefig(f"pic{i}.png")
                        # self.ax1.barh(categories, values, color=colormap(values))
                        # self.ax1.set_xlabel("Thể loại")
                        # self.ax1.set_ylabel("Số lượng")
                        # self.ax1.set_title("Top 10 số lượng theo thể loại")
                    else:
                        df_sorted_ = self.clean_data.sort_values(by=item[0])
                        df_ = pd.concat([df_sorted_.head(5), df_sorted_.tail(5)])
                        # print(df_)
                        # biểu đồ cột thì sẽ lấy 5 cột max và 5 cột min của dữ liệu số
                        text = df_[item[1]]
                        values = df_[item[0]]
                        var_axs[i].barh(text, values, color=colormap(values))
                        var_axs[i].set_xlabel(item[1])
                        var_axs[i].set_ylabel(item[0])
                        var_axs[i].set_title(f"Top 10")
                        var_canvas[i].draw()
                        var_fig[i].savefig(f"pic{i}.png")

                else:
                    # vẽ biểu đồ đường

                    var_axs[i].plot(self.clean_data.index, self.clean_data[item[0]], label=f"{item[0]}", linestyle='-', color='skyblue')
                    var_axs[i].plot(self.clean_data.index, self.clean_data[item[1]], label=f"{item[1]}", linestyle='--', color='orange')
                    # var_axs[i].plot(self.clean_data.index, self.clean_data[item[0]], label=f"{item[0]}", marker='o', linestyle='-', color='skyblue')
                    # var_axs[i].plot(self.clean_data.index, self.clean_data[item[1]], label=f"{item[1]}", marker='s', linestyle='--', color='orange')
                    var_axs[i].set_xlabel(item[0])
                    var_axs[i].set_ylabel(item[1])
                    var_axs[i].set_title(f"{item[0]} and {item[1]}")
                    var_axs[i].legend()
                    var_canvas[i].draw()
                    # self.ax2.plot(x_data, y_data1, label="Dữ liệu 1", marker='o', linestyle='-', color='skyblue')
                    # self.ax2.plot(x_data, y_data2, label="Dữ liệu 2", marker='s', linestyle='--', color='orange')

                    # self.ax2.set_xlabel('Thể loại')
                    # self.ax2.set_ylabel('Số lượng')
                    # self.ax2.set_title('Data 1 and Data 2')
                    # self.ax2.legend()
                    var_fig[i].savefig(f"pic{i}.png")
            pass

    
    def save_img_to_file(self):
        selected_value = self.figures.get()
        file_path = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG files", "*.png")])
        if file_path:
            if selected_value == "Figure 1":
                self.fig1.savefig(file_path)
                print("Lưu ảnh thành công!!!!")
                return
            elif selected_value == "Figure 2":
                self.fig2.savefig(file_path)
                print("Lưu ảnh thành công!!!!")
                return 
            elif selected_value == "Figure 3":
                self.fig3.savefig(file_path)
                print("Lưu ảnh thành công!!!!")
                return 
            elif selected_value == "Figure 4":
                self.fig4.savefig(file_path)
                print("Lưu ảnh thành công!!!!")
                return
            else:
                return
        else:
            print("Lưu ảnh không thành công!")   

    def switch_frames_back(self):
        # back về fram3 -> frame1
        # frame 2 -> frame 1
        if self.current_frame_l == self.f3_left and self.current_frame_r == self.f3_right:
            self.f3_left.place_forget()
            self.f3_right.place_forget()
            self.f1_left.place(x=0, y=0, width=400, height=700)
            self.f1_right.place(x=400, y=0, width=1000, height=700)
            self.current_frame_l = self.f1_left
            self.current_frame_r = self.f1_right
        if self.current_frame_l == self.f2_left and self.current_frame_r == self.f2_right:
            self.f2_left.place_forget()
            self.f2_right.place_forget()
            self.f1_left.place(x=0, y=0, width=400, height=700)
            self.f1_right.place(x=400, y=0, width=1000, height=700)
            self.current_frame_l = self.f1_left
            self.current_frame_r = self.f1_right

            
    def make_center(self):
        self.root.update_idletasks()
        width = self.root.winfo_width() # width of root
        height = self.root.winfo_height() # height of root

        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)

        self.root.geometry(f"{width}x{height}+{x}+{y}")
        return

    
    def export_word(self):
        # Tạo một tài liệu mới
        if self.analysis_data is not None:
            doc = Document()

            # Đặt kích thước giấy A4
            section = doc.sections[0]
            section.page_width = Cm(21) # Chiều rộng A4 (21 cm)
            section.page_height = Cm(29.7) # height 29.7 cm

            # Định dạng tiêu đề
            title = doc.add_paragraph("Báo cáo phân tích dữ liệu")
            title.runs[0].bold = True
            title.runs[0].font.size = Pt(18)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            body_text = [
                "I.Thông tin chung",
                self.analysis_data['recommend'],
                self.analysis_data['statistic'],
                "II. Kết quả xử lý số liệu"
            ]

            body_para = doc.add_paragraph(body_text[0])
            body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            body_para.paragraph_format.line_spacing = Pt(15) # dãn dòng 1.5
            body_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            body_para.runs[0].font.size = Pt(14)

            body_para = doc.add_paragraph(body_text[1])
            body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            body_para.paragraph_format.line_spacing = Pt(15) # dãn dòng 1.5
            body_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            body_para.runs[0].font.size = Pt(14)

            i = 0
            for index, row in self.analysis_data['statistic'].iterrows():
                print(row)
                text = f"{row.name}: mean = {row[1]}, mode = {row[2]}, median = {row[3]}, min = {row[4]}, max = {row[5]}"
                body_para = doc.add_paragraph(text)
                body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                body_para.paragraph_format.line_spacing = Pt(15) # dãn dòng 1.5
                body_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                body_para.runs[0].font.size = Pt(14)
                i += 1
            
            body_para = doc.add_paragraph(body_text[3])
            body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            body_para.paragraph_format.line_spacing = Pt(15) # dãn dòng 1.5
            body_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            body_para.runs[0].font.size = Pt(14)
            
            img0 = 'pic0.png'
            img1 = 'pic1.png'
            img2 = 'pic2.png'
            img3 = 'pic3.png'
            
            if os.path.exists(img0):
                doc.add_picture(img0)
            
            if os.path.exists(img1):
                doc.add_picture(img1)
            
            if os.path.exists(img2):
                doc.add_picture(img2)
            
            if os.path.exists(img3):
                doc.add_picture(img3)
            

            # Chèn hình ảnh
            

            end_text = [
                "III. Kết luận",
                "."*79,
                "."*79
            ]

            for i in end_text:
                end_para = doc.add_paragraph(i)
                end_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                end_para.paragraph_format.line_spacing = Pt(15) # dãn dòng 1.5
                end_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                end_para.runs[0].font.size = Pt(14)

            footer = "Ngày .... Tháng .... Năm .... \n\nKý tên\t\t"
            footer_doc = doc.add_paragraph(footer)
            footer_doc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            footer_doc.paragraph_format.line_spacing = Pt(15) # dãn dòng 1.5
            footer_doc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            footer_doc.runs[0].font.size = Pt(14)

            # Đặt font family cho toàn tài liệu
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"


            # Lưu tài liệu
            doc.save("Final_report.docx")
            print("Lưu thành công!!!")
            self.res_lbl.configure(text="Lưu thành công!")
            # self.res_lbl.place(x=200,y=650)
            # self.f3_left.after(2000, self.res_lbl.place_forget())
            messagebox.showinfo(title="Info", message="Lưu thành công!")
        else:
            messagebox.showerror(title="Error", message="Lỗi xuất dữ liệu do dữ liệu trống")
            # self.res_lbl.configure(text="Lưu thất bại!")
            # self.res_lbl.place(x=200,y=600)
            # self.f3_left.after(2000, self.res_lbl.place_forget())

    
    def submit_form(self):
        user_email = str(self.email_ent.get()).strip()
        subject_email = str(self.subject_ent.get()).strip()
        body_email = str(self.body_mail_text.get("1.0", "end-1c"))

        if not user_email or not subject_email or not body_email: # nếu 1 trong các trường trên trống
            messagebox.showerror(title='error', message="Vui lòng điền đầy đủ thông tin vào các trường trên!")
        else: # nếu đã điền đủ thông tin
            # kiểm tra tất cả các trường đã đầy đủ và không trống hay không
            if self.is_email(user_email):
                print("Valid")
                # email hợp lệ thì mới gửi email
                result = self.send_feedback(user_email, subject_email, body_email)
                self.result_lbl.configure(text=result)
                self.result_lbl.grid(row=5, column=2, padx=10, pady=10)
                self.result_lbl.after(2000, self.result_lbl.destroy) # sau 2s gửi thông báo cho user là có thành công hay không?
                self.lbl_frm_3.after(2000, self.clear_input) # sau 2s clear input
            else:
                # không hợp lệ thì thông báo lỗi
                print("Invalid")
                messagebox.showerror(title='error', message="Email không hợp lệ!!!")
            # print(user_email)

    
    def clear_input(self):
        self.email_ent.delete(0, "end")
        self.subject_ent.delete(0, "end")
        self.body_mail_text.delete("1.0", "end")
        return
    

    def is_email(self, email):
        """
        Hàm kiểm tra xem email có hợp lệ hay không bằng biểu thức chính quy
        """
        # r: raw string
        regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b'
        if re.fullmatch(regex, email):  # hàm này trả về None nếu mẫu không khớp
            return True
        else:
            return False
        
    
    # gửi email
    def send_feedback(self, user_mail, subject_mail, body_mail, attactment_file=None):
        """
        user_mail: email của user, sau khi nhà phát triển giải đáp các phản hồi của user thì sẽ lấy thông tin này để gửi
        subject_mail: tiêu đề phản hồi về ứng dụng của user
        body_mail: nội dung phản hồi
        attactment_file = None: tham số mặc định, user có thể đính kém file hoặc không?
        """
        # trong python quy ước hằng số viết in hoa và không thể gán lại
        SOURCE_EMAIL = 'nl142857nl@outlook.com.vn' # Email của ứng dụng
        PASSWORD = '@nl142857D1019'
        DESTINATION_EMAIL = 'rinnmusic2.2@gmail.com' # Email của nhóm phát triển
        MAIL_SERVER = 'smtp.office365.com'  # input: microsoft là smtp.office365.com ; google là smtp.gmail.com
        MAIL_PORT = 587 # microsort = 587, google 587 (bảo mật TLS - Transport Layer Security), 465 (SSL - Secure Socket Layer)
        subject_mail_ = subject_mail
        body_mail_ = "Xin chào nhóm phát triển ứng dụng,\r\n" + f"Feedback từ: {user_mail}" + "\r\n" + body_mail + "\r\n" + "Trân trọng!"

        email = EmailMessage()
        email['To'] = DESTINATION_EMAIL
        email['From'] = SOURCE_EMAIL
        email['Subject'] = subject_mail_
        email.set_content(body_mail_)

        if attactment_file: # nếu file tồn tại
            with open(file=attactment_file, mode='rb') as f:
                data = f.read()
                email.add_attachment(data, maintype='text', subtype='plain', filename=f.name)

        try:
            connection = smtplib.SMTP(host=MAIL_SERVER, port=MAIL_PORT)
            connection.starttls() # bảo mật tls: Transport Layer Security
            connection.login(SOURCE_EMAIL, PASSWORD)
            connection.send_message(from_addr=SOURCE_EMAIL, to_addrs=DESTINATION_EMAIL, msg=email)
            connection.quit()
            return "Gửi thành công! Bạn sẽ nhận được phản hồi trong vòng 48h!"
            
        except Exception as e:
            print(e)
            return f"{e}"

    ################################################
    # Trần Bá Quang
    ################################################


    ################################################
    # Trần Bá Quang
    ################################################



    ################################################
    # Bùi Nhật Minh
    ################################################




    ################################################
    # Bùi Nhật Minh
    ################################################



    ################################################
    # Nguyễn Quang Trường
    ################################################



    ################################################
    # Nguyễn Quang Trường
    ################################################


if __name__ == '__main__':
    app = App()
    app.root.mainloop()